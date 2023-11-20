
import docx

FOLDER_PATH = "/Users/everchicas/GIT/ReadPath/test/result"
NAME_FILE = "file"

wordFile = docx.Document()
wordFile.add_heading('Document',0)
data = (
    (1,"Valor 1"),
    (1,"Valor 2"),
    (1,"Valor 3"),
    (1,"Valor 4"),
    (1,"Valor 5")
)
table = wordFile.add_table(rows=1,cols=2)
row = table.rows[0].cells
row[0].text = 'Id'
row[1].text = 'Name'

for id, name in data:
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name

wordFile.save(f'{FOLDER_PATH}/{NAME_FILE}.docx')